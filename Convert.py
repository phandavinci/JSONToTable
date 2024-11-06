import pandas as pd
from openpyxl import Workbook
import docx
from openpyxl.styles import Font, Border, Side
import os
from docx.shared import RGBColor, Pt

def ToExcel(parsedJSON):
    if os.path.exists('output.xlsx'): os.remove('output.xlsx')
    wb = Workbook()
    ws = wb.active
    rows=1
    headers = ["Field Name", "Type"]
    tablename_font = Font(size=16, color='0f4761')
    header_font = Font(size=12, color='FF0000', bold=True)
    row1_font = Font(size=12, color='000000', bold=True)
    row2_font = Font(size=12, color="00339b", bold=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # for table_name, table_data in parsedJSON.items():
    #     ws.append([table_name])
    #     ws.cell(row=row, column=1).font = tablename_font
    #     row+=1
    ws.append(headers)
    ws.cell(row=rows, column=1).font = header_font
    ws.cell(row=rows, column=2).font = header_font
    rows+=1
    for rowitems in parsedJSON:
        for field_name, field_type in rowitems.items():
            ws.append([field_name, field_type])
            ws.cell(row=rows, column=1).font = row1_font
            ws.cell(row=rows, column=2).font = row2_font
            rows+=1
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = border
    wb.save("output.xlsx")
    return ws

    # with pd.ExcelWriter('output.xlsx', engine='openpyxl') as writer:
    #     start_row = 0
    #     for table_name, table_data in parsedJSON.items():
    #         rows = []
    #         df_header = pd.DataFrame({table_name: []})
    #         df_header.to_excel(writer, sheet_name='Sheet1', startrow=start_row, index=False)
    #         start_row+=1
    #         for row in table_data:
    #             for field_name, field_type in row.items():
    #                 rows.append({"Field Name": field_name, "Type": field_type})
    #         df_rows = pd.DataFrame(rows)
    #         df_rows.to_excel(writer, sheet_name='Sheet1', startrow=start_row, index=False)
    #         start_row+= len(df_rows)+2

def ToWord(JSON):
    if os.path.exists('output.docx'): os.remove('output.docx')
    print("In ToWord") #############
    abbrevation = set(['String', 'Boolean', 'Int', 'List'])
    header_size = col1_size = col2_size = Pt(12)
    font_style_common = 'Calibri'
    title_size = Pt(18)
    reqres_size = Pt(16)
    tablename_size = Pt(14)
    header_color = RGBColor(255, 0, 0)
    col1_color = RGBColor(0,0,0)
    col2_color = RGBColor(0, 51, 155)
    col2_link_color = RGBColor(26, 0, 225)

    doc = docx.Document()
    for title in JSON:
        run = doc.add_heading("", 2).add_run(title)
        run.font.bold = False
        run.font.size = title_size
        run.font.name = font_style_common
        for reqres in JSON[title]:
            run = doc.add_heading("", 3).add_run(reqres)
            run.font.bold = False
            run.font.size = reqres_size
            run.font.name = font_style_common
            for tablename in JSON[title][reqres]:
                if tablename != '': 
                    run = doc.add_heading("", 4).add_run(tablename)
                    run.font.bold = False
                    run.font.size = tablename_size
                    run.font.name = font_style_common
                excel = ToExcel(JSON[title][reqres][tablename])
                table = doc.add_table(rows=excel.max_row, cols=excel.max_column)
                table.style = 'Table Grid'
                for i, row in enumerate(excel.iter_rows(values_only=True)):
                    for j, col in enumerate(row):
                        cell = table.cell(i, j)
                        if col is not None:
                            cell.text = str(col)
                            cell_run = cell.paragraphs[0].runs[0]
                            cell_run.font.name = font_style_common
                            if i==0:
                                cell_run.font.bold = True
                                cell_run.font.color.rgb = header_color
                                cell_run.font.size = header_size
                            elif j==0:
                                cell_run.font.bold = True
                                cell_run.font.color.rgb = col1_color
                                cell_run.font.size = col1_size
                            elif j==1:
                                if str(col) in abbrevation:
                                    cell_run.font.bold = True
                                    cell_run.font.color.rgb = col2_color
                                    cell_run.font.size = col2_size
                                else:
                                    cell_run.font.bold = True
                                    cell_run.font.italic = True
                                    cell_run.font.underline = True
                                    cell_run.font.color.rgb = col2_link_color
                                    cell_run.font.size = col2_size
                        else:
                            cell.text = ""
        print(title, "done successfully")
    doc.save('output.docx')
